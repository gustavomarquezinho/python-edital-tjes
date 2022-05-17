"""
-> Módulo que habilita realizar modificações em documentos.
"""


from docx import Document
import docx.text.paragraph

from docx.opc.exceptions import \
    PackageNotFoundError

from tkinter.filedialog import asksaveasfilename
from .gui import show_message


class Docs:
    """
    -> Classe principal e única contendo todas as funções e métodos.
        :author: Gustavo Marquezinho
    """

    def __init__(self, path: str):
        """
        -> Inicia a classe e carrega o arquivo.
            :param path: caminho do arquivo.
        """

        self.finded, self.changed = 0, 0

        self.path = path
        self.loaded = None

    ######

    def load_file(self) -> bool:
        """
        -> Carrega o arquivo a ser alterado.
            :return: sucesso ou falha
        """

        try:
            self.loaded = Document(self.path)

        except PackageNotFoundError:
            show_message('(Erro) Arquivo ou diretório do arquivo base não encontrado!', True)
            return False

        except PermissionError:
            show_message('(Erro) Permissão negada para a leitura do arquivo base!', True)
            return False
        return True

    ######

    def loop_docs(self, search: dict) -> None:
        """
        -> Realiza um laço pelos parágrafos e tabelas.
            :param search: dicionário contendo os valores.
            :return: None
        """

        for p in self.loaded.paragraphs:
            self.search_paragraph(p, search)

        for table in self.loaded.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.search_paragraph(p, search)

    ######

    def search_paragraph(self, paragraph: docx.text.paragraph.Paragraph, search: dict) -> None:
        """
        -> Procura pelas incidências de certos valores.
            :param paragraph: parágrafo onde será procurado.
            :param search: dicionário contendo os valores.
            :return: None
        """

        if '#' in paragraph.text:
            self.finded += 1

            for key in search:
                if key in paragraph.text:
                    self.replace_paragraph(paragraph, key, search[key])

    ######

    def replace_paragraph(self, p, key: str, value: str) -> None:
        """
        -> Substitui o texto de um parágrafo mantendo o estilo.
            :param p: parágrafo a ser alterado.
            :param key: antigo valor.
            :param value: novo valor.
            :return: None
        """

        if value == '':
            paragraph = p._element
            
            paragraph.getparent().remove(paragraph)
            p.paragraph = p._element = None

            self.changed += 1
            return

        inline = p.runs

        for i in range(len(inline)):
            if key in inline[i].text:
                inline[i].text = inline[i].text.replace(key, value)
                self.changed += 1

    ######

    def save_file(self, name: str) -> bool:
        """
        -> Salva o arquivo onde o usuário desejar e apresenta as informações sobre a escrita.
            :param name: nome do arquivo.
            :return: Resultado da tentativa
        """

        filetype = [
            ('Documento do Microsoft Word', '*.docx')
        ]
        
        path = asksaveasfilename(filetypes=filetype, defaultextension=filetype[0][1], initialfile=name)

        if len(path) == 0:
            return False 

        try:
            self.loaded.save(path)

        except FileNotFoundError:
            show_message('(Erro) Arquivo ou diretório de destino não encontrado!')
            return False

        except PermissionError:
            show_message('(Erro) Permissão negada para o salvamento do arquivo!')
            return False

        show_message(F'Documento gerado com sucesso ({self.changed} campos alterados de {self.finded} encontrados).')

        self.finded, self.changed = 0, 0
        return True
